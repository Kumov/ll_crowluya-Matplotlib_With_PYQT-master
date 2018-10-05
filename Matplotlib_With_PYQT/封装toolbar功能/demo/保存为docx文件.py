from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

document = Document()

# 设置属性
document.styles['Normal'].font.name = u'黑体'
document.add_heading('This is my title 我的标题', 0)

# #添加文本
# paragraph = document.add_paragraph(u'添加了文本')
# #设置字号
# run = paragraph.add_run(u'设置字号')
# run.font.size=Pt(24)
#
# #设置字体
# run = paragraph.add_run('Set Font,')
# run.font.name='Consolas'
#
# #设置中文字体
# run = paragraph.add_run(u'设置中文字体，')
# run.font.name=u'宋体'
# r = run._element
# r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#
# #设置斜体
# run = paragraph.add_run(u'斜体、')
# run.italic = True

#设置粗体
# run = paragraph.add_run(u'粗体').bold = True

pra = '油藏描述，简称RDS技术服务（ReservoirDescriptionService），就是对油藏各种特征进行三维空间的定量描述和表征以至预测。' \
      '以沉积学、构造地质学、储层地质学和石油地质学的理论为指导，综合应用地震地层学、测井地质学、储层地质学、地质统计学等方法，' \
      '以计算机为手段，定性与定量地描述三维空间中的油气藏类型、外部几何形状、规模大小、油藏内部结构、储层参数变化和流体性质及分布等，' \
      '揭示油气在三维空间中的分布与变化规律。'
p = document.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(pra)
run.font.name = u'宋体'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.color.rgb = RGBColor(54, 95, 145)
run.font.size = Pt(15)

# 居中
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 首行缩进
p.paragraph_format.first_line_indent = Inches(0.3)
p.paragraph_format.first_line_indent = Inches(0.3)
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.line_spacing = Pt(18)


p = document.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(u'所有曲线图像 ')
run.font.color.rgb = RGBColor(54, 95, 145)
run.font.size = Pt(36)

# 添加图片
pic = document.add_picture('Img/image-1.png', width=Inches(6))
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置

pic = document.add_picture('Img/image-2.png', width=Inches(6))
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置

pic = document.add_picture('Img/image-3.png', width=Inches(6))
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置

# 添加表格
rows = 3
cols = 3
table = document.add_table(rows=rows, cols=cols, style="Table Grid")  # 添加2行3列的表格

for i in range(rows):
    tr = table.rows[i]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), "450")
    trPr.append(trHeight)  # 表格高度设置
# table.autofit = False
col = table.columns[1]
col.width = Inches(5)
arr = [u'序号', u"类型", u"详细描述"]
heading_cells = table.rows[0].cells
for i in range(cols):
    p = heading_cells[i].paragraphs[0]
    run = p.add_run(arr[i])
    run.font.color.rgb = RGBColor(54, 95, 145)  # 颜色设置，这里是用RGB颜色
    run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

#
td1 = 'td=0.12'
td2 = 'td=0.144'
# td3 = 'td=0.176'
#
table.cell(-1, 0).text = td1
table.cell(2, 0).text = td2
# table.cell(3, 0).text = td3
#
# # 添加描述
table.cell(1, 1).text = '添加类型'
table.cell(1, 2).text = '添加描述'


table.add_row()
document.save('data/test3.docx')
